"""
generate_report_docx.py — Generates a formatted project report DOCX.
Run from the repo root: python scripts/generate_report_docx.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT_PATH = Path(__file__).resolve().parent.parent / "NCIIPC_Project_Report.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, bold=False, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(shading)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "1F3A5F")
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
            if r_idx % 2 == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), "EFF4FB")
                row.cells[c_idx]._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NCIIPC Compliance Scoring Pipeline")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x0F, 0x1F, 0x38)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Project Report")
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x5A, 0x6A, 0x7E)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n")
    meta.add_run("Stack: Python 3.13 · Flask · C++17 (pybind11) · OpenRouter LLMs\n")
    meta.add_run("Scope: NCIIPC Controls PC1, PC2, PC3\n")

    doc.add_page_break()

    # ── 1. Abstract ───────────────────────────────────────────────────────
    add_heading(doc, "1. Abstract", 1)
    add_para(doc,
        "This report describes the design, implementation, and evaluation of an automated "
        "cybersecurity compliance auditing system for organisations operating Critical "
        "Information Infrastructure (CII) under India's National Critical Information "
        "Infrastructure Protection Centre (NCIIPC) framework. The system accepts PDF, DOCX, "
        "TXT, or URL inputs, extracts and indexes document text using a high-performance C++ "
        "BM25+ retrieval engine, and scores each document against three compliance controls "
        "(PC1, PC2, PC3) using an ensemble of three large language models (LLMs) with a "
        "regex-based fallback. Scores are aggregated into per-control maturity levels (L1–L5) "
        "and an overall weighted company score, accompanied by natural-language audit findings "
        "generated by an LLM. A smoke test suite of nine synthetic documents validates all "
        "pipeline stages, achieving 9/9 on the scorer."
    )

    # ── 2. Problem Statement ──────────────────────────────────────────────
    add_heading(doc, "2. Problem Statement and Motivation", 1)
    add_para(doc,
        "Indian organisations designated as Critical Information Infrastructure operators are "
        "required to demonstrate compliance with NCIIPC controls. Manual auditing of compliance "
        "documents is slow, inconsistent, and expert-intensive. A single audit cycle may involve "
        "reviewing dozens of policy documents, annual reports, SOC charters, and asset registers "
        "spanning hundreds of pages."
    )
    add_para(doc,
        "The specific problems this system addresses are:"
    )
    add_bullet(doc, "Auditors must manually read and cross-reference large document sets to assess compliance.")
    add_bullet(doc, "Scoring is subjective and varies between reviewers.")
    add_bullet(doc, "There is no machine-readable compliance signal for CII-specific controls (PC1–PC3).")
    add_bullet(doc, "Generating structured audit findings from raw document evidence is time-consuming.")
    add_para(doc,
        "This system automates the evidence-retrieval, scoring, and finding-generation steps, "
        "reducing a multi-day manual audit to a process that completes in under five minutes per document set."
    )

    # ── 3. NCIIPC Controls ────────────────────────────────────────────────
    add_heading(doc, "3. NCIIPC Controls Covered", 1)
    add_para(doc,
        "The system evaluates three NCIIPC Protective Controls. Each control maps to a set of "
        "compliance sub-fields assessed individually by BM25 retrieval and LLM voting."
    )
    add_table(doc,
        ["Control", "Name", "What It Checks", "Sub-fields"],
        [
            ("PC1", "CII Identification",
             "Existence and quality of a CII asset register, use of six classification criteria, periodic review",
             "register, criteria, review, weakness, strength"),
            ("PC2", "Interdependency Mapping",
             "Vertical and horizontal dependency maps, inbound/outbound mapping, SLA documentation",
             "vertical, horizontal, weakness, strength"),
            ("PC3", "Cybersecurity Governance",
             "CISO reporting line, ISD establishment, 24/7 SOC, audit independence, ISD four functions",
             "ciso, isd, soc, audit, functions, weakness"),
        ]
    )

    # ── 4. System Architecture ────────────────────────────────────────────
    add_heading(doc, "4. System Architecture", 1)
    add_para(doc,
        "The system is composed of four layers: a web interface, a Python orchestration layer, "
        "a C++ performance module, and external LLM APIs."
    )

    add_heading(doc, "4.1 Component Overview", 2)
    add_table(doc,
        ["Component", "Technology", "Role"],
        [
            ("Web Frontend",      "HTML + CSS + JavaScript (index.html)",     "Single-page UI: file upload, URL input, score display"),
            ("Flask Server",      "Python 3.13 + Flask 3.0 (app.py)",          "HTTP endpoint, rate limiting, temp file management"),
            ("Text Extraction",   "pdfplumber, python-docx, trafilatura (pipeline.py)", "Multi-format text extraction from PDF/DOCX/TXT/URL"),
            ("BM25 Engine",       "C++17 + pybind11 (nciipc-prep/)",           "Sliding-window chunking + BM25+ inverted index"),
            ("Scoring Engine",    "Python (scorer.py)",                        "Field retrieval, LLM voting orchestration, aggregation"),
            ("LLM Voting",        "OpenRouter API via openai SDK (explain.py)","3-model ensemble voting: COMPLIANT/PARTIAL/NON_COMPLIANT"),
            ("LF Fallback",       "Python regex (labeling_functions.py)",      "20 regex labeling functions, Snorkel-style fallback"),
            ("Findings Generator","OpenRouter + Ollama fallback (explain.py)", "3-sentence audit finding per control"),
        ]
    )

    add_heading(doc, "4.2 End-to-End Pipeline", 2)
    add_para(doc, "The pipeline processes documents in eight stages:")
    steps = [
        ("Text Extraction",    "PDF/DOCX/TXT/URL → plain text via pdfplumber, python-docx, or trafilatura. "
                               "Repeated lines (headers/footers) are stripped from PDFs using frequency counting."),
        ("Chunking",           "Text is tokenised by whitespace and split into 512-token sliding windows "
                               "with a 384-token stride (128-token overlap). Implemented in C++ for speed."),
        ("BM25 Indexing",      "An inverted index is built over all chunks using BM25+ (k₁=1.5, b=0.75, δ=1.0). "
                               "The C++ implementation uses partial_sort for efficient top-k retrieval."),
        ("Field Retrieval",    "For each of 16 compliance sub-fields, a keyword query retrieves the top-5 chunks "
                               "with BM25 score ≥ 1.0. Chunks below the threshold are discarded."),
        ("LLM Ensemble Voting","Each retrieved chunk is sent to three free LLMs on OpenRouter. Each model votes "
                               "COMPLIANT (1.0), PARTIAL (0.5), or NON_COMPLIANT (0.0). The mean vote is the chunk score."),
        ("LF Fallback",        "If all LLM calls fail, 6–8 regex labeling functions score the chunk instead. "
                               "Non-firing functions abstain; the mean of firing functions is the chunk score. "
                               "If all functions abstain, the score is 0.0 (absence of evidence = non-compliant)."),
        ("Score Aggregation",  "Field score = mean(chunk scores) × 100. Control score = mean(field scores). "
                               "Company score = (PC1×1.00 + PC2×0.90 + PC3×0.95) / 2.85."),
        ("Finding Generation", "An LLM generates a 3-sentence audit finding per control, grounded in weak, "
                               "medium, and strong evidence snippets from the field results."),
    ]
    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f"[{i}] {title}: ")
        run_num.bold = True
        p.add_run(desc)

    add_heading(doc, "4.3 Call Graph", 2)
    for line in [
        "app.py /score",
        "  ├── pipeline.extract_text()       [PDF/DOCX/TXT/URL → text]",
        "  ├── scorer.score_documents(texts)",
        "  │     ├── nciipc_cpp.build_index()   [C++: chunk + index]",
        "  │     └── for each of 16 fields:",
        "  │           ├── nciipc_cpp.query_bm25()   [top-5 chunks]",
        "  │           ├── explain.vote_chunk()       [3 LLMs via OpenRouter]",
        "  │           └── labeling_functions.score_chunk()  [regex fallback]",
        "  └── explain.explain_control()     [audit finding per control]",
        "        ├── OpenRouter (4 retries + jitter)",
        "        └── Ollama local (mistral fallback)",
    ]:
        add_code(doc, line)

    # ── 5. Algorithms ─────────────────────────────────────────────────────
    add_heading(doc, "5. Algorithms and Techniques", 1)

    add_heading(doc, "5.1 Sliding-Window Chunking", 2)
    add_para(doc,
        "Documents are tokenised by whitespace and divided into overlapping chunks to ensure "
        "evidence spanning paragraph boundaries is captured. Parameters:"
    )
    add_bullet(doc, "Window size: 512 tokens")
    add_bullet(doc, "Stride: 384 tokens")
    add_bullet(doc, "Overlap: 128 tokens")
    add_para(doc,
        "A 1,500-token document produces 4 chunks. The 128-token overlap ensures no compliance "
        "signal shorter than 128 tokens is split across two non-overlapping windows."
    )

    add_heading(doc, "5.2 BM25+ Retrieval", 2)
    add_para(doc,
        "BM25+ (Lv & Zhai, 2011) extends Okapi BM25 with a δ parameter that guarantees a "
        "positive score contribution from every matching term, preventing very long documents "
        "from being unfairly penalised. The scoring formula for a query Q and chunk D is:"
    )
    add_code(doc, "score(Q,D) = Σ_t  idf(t) · [ tf(t,D)·(k₁+1) / (tf(t,D) + k₁·(1−b+b·|D|/avgdl)) + δ ]")
    add_para(doc, "Where: k₁=1.5  b=0.75  δ=1.0  idf(t)=ln((N−df+0.5)/(df+0.5)+1)")
    add_para(doc,
        "The implementation (bm25.cpp) builds an in-memory inverted index and uses C++ "
        "partial_sort for efficient top-k retrieval without sorting the entire result set."
    )

    add_heading(doc, "5.3 LLM Ensemble Voting", 2)
    add_para(doc,
        "Each retrieved chunk is evaluated by three independently-prompted LLMs. Ensemble voting "
        "reduces single-model bias and hallucination risk:"
    )
    add_bullet(doc, "meta-llama/llama-3.1-8b-instruct:free")
    add_bullet(doc, "google/gemma-2-9b-it:free")
    add_bullet(doc, "mistralai/mistral-small-24b-instruct-2501:free")
    add_para(doc,
        "Each model is prompted with the compliance field description and chunk text, and asked "
        "to respond with exactly one word: COMPLIANT, PARTIAL, or NON_COMPLIANT. Temperature is "
        "set to 0.0 and max_tokens to 10 for determinism. The chunk score is the arithmetic mean "
        "of all successful votes. The field result includes an llm_coverage metric (0.0–1.0) "
        "indicating what fraction of chunks were LLM-scored vs. regex-scored."
    )

    add_heading(doc, "5.4 Regex Labeling Functions (Snorkel-style)", 2)
    add_para(doc,
        "Inspired by the Snorkel weak-supervision framework, 20 labeling functions (LFs) provide "
        "a deterministic fallback when LLMs are unavailable. Each LF is a narrow regex detector "
        "that returns COMPLIANT (1.0), PARTIAL (0.5), NON_COMPLIANT (0.0), or ABSTAIN (None). "
        "The chunk score is the mean of non-abstaining votes. If all LFs abstain, the score is "
        "0.0 — absence of evidence is treated as non-compliance, not uncertainty."
    )
    add_table(doc,
        ["Control", "LF", "Signal detected"],
        [
            ("PC1", "lf_pc1_register_exists",  "Phrase 'CII asset register' or 'critical infrastructure list'"),
            ("PC1", "lf_pc1_six_criteria",     "Count of the 6 NCIIPC criteria keywords (≥5 → COMPLIANT, ≥3 → PARTIAL)"),
            ("PC1", "lf_pc1_no_register",      "Negation near 'CII register' → NON_COMPLIANT"),
            ("PC1", "lf_pc1_scada_excluded",   "SCADA/OT explicitly excluded from scope → NON_COMPLIANT"),
            ("PC2", "lf_pc2_both_maps",        "Both vertical and horizontal interdependency mentioned"),
            ("PC2", "lf_pc2_inbound_outbound", "Both inbound and outbound dependencies mentioned"),
            ("PC3", "lf_pc3_ciso_reports_head","CISO reports to head/CEO/board → COMPLIANT"),
            ("PC3", "lf_pc3_soc_24x7",         "SOC + 24/7 mention (guards against negation to avoid double-counting)"),
            ("PC3", "lf_pc3_four_functions",   "Count of ISD four functions (≥3 → COMPLIANT, ≥2 → PARTIAL)"),
        ]
    )

    add_heading(doc, "5.5 Score Aggregation and Maturity Mapping", 2)
    add_table(doc,
        ["Level", "Label", "Score range", "Interpretation"],
        [
            ("L5", "Optimising",  "≥ 80", "Processes continuously improved; metrics-driven"),
            ("L4", "Managed",     "≥ 65", "Processes measured and controlled"),
            ("L3", "Defined",     "≥ 50", "Processes documented and standardised"),
            ("L2", "Developing",  "≥ 35", "Ad hoc structure; partial documentation"),
            ("L1", "Initial",     "< 35",  "Unpredictable; no formal process"),
        ]
    )
    add_para(doc, "Weighted company score formula:")
    add_code(doc, "company_score = (PC1×1.00 + PC2×0.90 + PC3×0.95) / 2.85")
    add_para(doc,
        "PC1 carries the highest weight (1.00) as CII identification is the foundation for all "
        "other controls. PC3 (governance) is weighted at 0.95, and PC2 (interdependency) at 0.90."
    )

    # ── 6. Design Rationale ───────────────────────────────────────────────
    add_heading(doc, "6. Design Rationale", 1)
    add_table(doc,
        ["Decision", "Alternatives considered", "Rationale"],
        [
            ("BM25+ over dense (vector) retrieval",
             "FAISS + sentence-transformers",
             "BM25 is interpretable, requires no GPU, has no embedding model to maintain, "
             "and performs well on keyword-heavy compliance text. Dense retrieval adds "
             "significant infrastructure cost for marginal gain on this domain."),
            ("C++ BM25 implementation",
             "Python rank_bm25 library",
             "Indexing 50 MB of PDFs produces tens of thousands of chunks. C++ partial_sort "
             "and unordered_map give 10–50× faster build and query times vs. pure Python."),
            ("3-model LLM ensemble",
             "Single model, GPT-4",
             "Free-tier models have high error rates individually. Averaging three models "
             "reduces variance and cost. Ensemble agrees on clear cases; disagreement "
             "indicates genuine ambiguity (→ PARTIAL)."),
            ("Regex LF fallback",
             "Always require LLMs",
             "OpenRouter free models have rate limits and outages. A deterministic fallback "
             "ensures the system always produces a score, with llm_coverage indicating "
             "how much of the result is LLM-backed."),
            ("threaded=False in Flask",
             "Gunicorn / multi-worker",
             "The C++ module uses global state (g_index, g_chunks) with no mutex. "
             "Single-threaded mode is safe for a local/demo deployment. Production use "
             "would require per-request index isolation."),
            ("Score-band evidence (weak/medium/strong)",
             "Top-N fields only",
             "Fields scoring 50–64 were previously excluded from the LLM finding prompt, "
             "leaving the model with no grounding for a third of possible cases. "
             "Three bands ensure all evidence reaches the prompt."),
        ]
    )

    # ── 7. Implementation Details ─────────────────────────────────────────
    add_heading(doc, "7. Implementation Details", 1)

    add_heading(doc, "7.1 Project Structure", 2)
    for line in [
        "nciipc/",
        "├── app.py                    Flask server (HTTP, rate limit, temp files)",
        "├── requirements.txt          86 Python dependencies",
        "├── .env                      API key + model config (git-ignored)",
        "├── .gitignore                Excludes .env, .venv, data/, logs/, build/",
        "├── scripts/",
        "│   ├── pipeline.py           Text extraction + CLI entry point",
        "│   ├── scorer.py             BM25 retrieval + LLM voting + aggregation",
        "│   ├── explain.py            LLM voting + audit finding generation",
        "│   ├── labeling_functions.py 20 regex compliance detectors",
        "│   └── smoke_test.py         9-document end-to-end test suite",
        "├── nciipc-prep/              C++ pybind11 module",
        "│   ├── src/chunker.cpp       Sliding-window tokeniser",
        "│   ├── src/bm25.cpp          BM25+ inverted index",
        "│   └── src/bindings.cpp      pybind11 Python bindings",
        "└── web/index.html            Single-page frontend (no build step)",
    ]:
        add_code(doc, line)

    add_heading(doc, "7.2 Reliability and Error Handling", 2)
    add_bullet(doc, "LLM calls use exponential backoff with ±25% jitter: wait ≈ 1s, 2s, 4s across 4 attempts per model.")
    add_bullet(doc, "If all OpenRouter retries fail for explain_control(), the system falls back to local Ollama (model: mistral).")
    add_bullet(doc, "If all LLMs fail for vote_chunk(), scoring falls back to regex labeling functions.")
    add_bullet(doc, "If no BM25 chunks score ≥ 1.0 for a field, the field scores 0.0 (non-compliant by absence).")
    add_bullet(doc, "All LLM calls (successes and failures) are logged to logs/llm_calls/ as daily JSONL files.")
    add_bullet(doc, "Uploaded files are written to OS temp and deleted in a finally block regardless of outcome.")

    add_heading(doc, "7.3 Security Measures", 2)
    add_table(doc,
        ["Measure", "Implementation"],
        [
            ("Upload size limit",    "50 MB hard cap via Flask MAX_CONTENT_LENGTH → HTTP 413"),
            ("File type validation", "Extension allowlist: .pdf, .docx, .txt, .text → HTTP 400 otherwise"),
            ("Rate limiting",        "10 requests / 60 s per IP; idle IPs evicted to prevent memory growth"),
            ("Temp file cleanup",    "os.unlink() in finally block; no uploaded data persists after response"),
            ("Thread safety",        "threaded=False prevents concurrent writes to C++ global index state"),
            ("API key protection",   ".env loaded at runtime; file is git-ignored and never hardcoded"),
        ]
    )

    # ── 8. Evaluation ─────────────────────────────────────────────────────
    add_heading(doc, "8. Evaluation", 1)

    add_heading(doc, "8.1 Smoke Test Results", 2)
    add_para(doc,
        "The smoke test validates all pipeline stages using 9 purpose-written synthetic documents "
        "(3 per control × 3 compliance levels: COMPLIANT, PARTIAL, NON_COMPLIANT). Each document "
        "is ~350 words of realistic compliance text."
    )
    add_table(doc,
        ["Stage", "Result", "Notes"],
        [
            ("Labeling Functions", "9 / 9  PASS", "All 9 docs produce scores in the expected range"),
            ("BM25 Retrieval",     "3 / 3  PASS", "Each control query returns ≥1 result with score > 0"),
            ("Scorer (end-to-end)","9 / 9  PASS", "All maturity labels match expected COMPLIANT/PARTIAL/NON_COMPLIANT"),
            ("Explain (Ollama)",   "PASS",         "3-sentence findings generated; each ≥ 50 characters"),
        ]
    )
    add_para(doc,
        "Note: The scorer previously failed 3/9 cases (pc1_partial, pc2_partial, pc3_partial scored "
        "as NON_COMPLIANT). Root causes identified and fixed: (1) OPENROUTER_VOTING_MODELS in .env "
        "contained unavailable model IDs, silently forcing LF-only scoring; (2) score_chunk returned "
        "0.5 for all-abstain cases, inflating NON_COMPLIANT documents but deflating PARTIAL ones when "
        "combined with NON_COMPLIANT LF signals."
    )

    add_heading(doc, "8.2 Sample Output (Archived Test Run)", 2)
    add_para(doc,
        "The following scores were produced from a real test run on power-sector compliance "
        "documents (archived in archive/data/raw/):"
    )
    add_table(doc,
        ["Control", "Score", "Maturity", "Interpretation"],
        [
            ("PC1 — CII Identification",     "37.5 / 100", "L2 Developing", "Register exists but incomplete criteria coverage"),
            ("PC2 — Interdependency Mapping", "39.2 / 100", "L2 Developing", "Vertical mapping present; horizontal gaps"),
            ("PC3 — Cybersecurity Governance","25.8 / 100", "L1 Initial",    "No evidence of 24/7 SOC or CISO at board level"),
            ("Company (weighted)",            "34.1 / 100", "L1 Initial",    "Significant compliance gaps across all controls"),
        ]
    )

    # ── 9. Limitations ────────────────────────────────────────────────────
    add_heading(doc, "9. Limitations", 1)
    add_bullet(doc,
        "Single-threaded server: Flask runs with threaded=False due to global C++ state. "
        "The system cannot handle concurrent requests; production deployment requires refactoring "
        "the C++ module to use per-request index instances.")
    add_bullet(doc,
        "Free-tier LLM reliability: OpenRouter free models have rate limits, outages, and "
        "occasional model removals. The regex fallback mitigates this but produces lower-quality scores.")
    add_bullet(doc,
        "No authentication: The /score endpoint is unauthenticated. Anyone with network access "
        "can submit requests (rate limiting provides only partial protection).")
    add_bullet(doc,
        "Scope limited to PC1–PC3: NCIIPC defines additional controls (PC4–PC9). The current "
        "BM25 queries and labeling functions cover only the first three.")
    add_bullet(doc,
        "English-only: Text extraction and LLM prompting assume English-language documents. "
        "Hindi or bilingual documents will produce degraded results.")
    add_bullet(doc,
        "No result persistence: All scoring is ephemeral. There is no database, so historical "
        "audit results cannot be retrieved or trended without re-running the pipeline.")
    add_bullet(doc,
        "BM25 threshold is fixed: MIN_BM25_SCORE = 1.0 is hardcoded and was not empirically "
        "validated against a labelled holdout set.")

    # ── 10. Future Work ───────────────────────────────────────────────────
    add_heading(doc, "10. Future Work", 1)
    add_table(doc,
        ["Priority", "Enhancement", "Description"],
        [
            ("High",   "Concurrent request support",
             "Refactor C++ bindings to use per-request index instances (eliminate global state). "
             "Switch to gunicorn or uvicorn for production serving."),
            ("High",   "PC4–PC9 coverage",
             "Add BM25 queries and labeling functions for the remaining NCIIPC controls."),
            ("Medium", "Database persistence",
             "Store audit results in PostgreSQL to enable compliance trend analysis over time."),
            ("Medium", "Validated thresholds",
             "Calibrate BM25 score threshold and maturity boundaries against a labelled holdout set "
             "of real CII compliance documents."),
            ("Medium", "Docker packaging",
             "Containerise the application (including C++ build) for reproducible deployment."),
            ("Low",    "LLM finding caching",
             "Cache explain_control() outputs by (control, score_range, document_hash) to reduce "
             "redundant API calls when the same document is re-submitted."),
            ("Low",    "Log rotation",
             "Implement automatic compression/archival of JSONL logs older than 90 days."),
        ]
    )

    # ── 11. References ────────────────────────────────────────────────────
    add_heading(doc, "11. References", 1)
    refs = [
        "NCIIPC, Government of India. Guidelines for Protection of Critical Information Infrastructure. "
        "National Critical Information Infrastructure Protection Centre, 2021.",

        "Lv, Y. & Zhai, C. (2011). Lower-Bounding Term Frequency Normalization. "
        "Proceedings of the 20th ACM CIKM, pp. 7–16.",

        "Ratner, A., Bach, S. H., Ehrenberg, H., Fries, J., Wu, S., & Ré, C. (2017). "
        "Snorkel: Rapid Training Data Creation with Weak Supervision. VLDB, 11(3), pp. 269–282.",

        "Robertson, S. & Zaragoza, H. (2009). The Probabilistic Relevance Framework: BM25 and Beyond. "
        "Foundations and Trends in Information Retrieval, 3(4), pp. 333–389.",

        "pdfplumber. https://github.com/jsvine/pdfplumber",

        "trafilatura. https://trafilatura.readthedocs.io/",

        "OpenRouter. https://openrouter.ai — Multi-model LLM API gateway.",

        "pybind11. https://pybind11.readthedocs.io — Seamless C++ and Python interoperability.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f"[{i}]  ")
        run_num.bold = True
        p.add_run(ref)

    # ── 12. Conclusion ────────────────────────────────────────────────────
    add_heading(doc, "12. Conclusion", 1)
    add_para(doc,
        "This project delivers a functional, end-to-end automated compliance auditor for the "
        "NCIIPC PC1–PC3 controls. The combination of a high-performance C++ BM25+ retrieval "
        "engine, a three-model LLM ensemble, and a deterministic regex fallback produces "
        "interpretable compliance scores with natural-language audit findings in under five "
        "minutes per document set. All nine smoke test cases pass, including the three PARTIAL "
        "compliance scenarios that were previously misclassified following bug fixes to the "
        "voting model configuration and the all-abstain fallback scoring logic."
    )
    add_para(doc,
        "The system is ready for use with real CII compliance documents and is structured for "
        "extension to additional NCIIPC controls. Key next steps are concurrent request support, "
        "coverage of PC4–PC9, and empirical calibration of scoring thresholds against a labelled "
        "holdout corpus."
    )

    doc.save(OUT_PATH)
    print(f"Report saved -> {OUT_PATH}")


if __name__ == "__main__":
    build()
