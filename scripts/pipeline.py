"""
pipeline.py — Single entry point for the NCIIPC compliance scoring pipeline.

Accepts: PDF, DOCX, TXT file paths, or a website URL.

Usage:
    python scripts/pipeline.py path/to/report.pdf
    python scripts/pipeline.py path/to/policy.docx
    python scripts/pipeline.py https://company.example.com/annual-report

Requires:  OPENROUTER_API_KEY set in a .env file
"""

import sys
import json
from pathlib import Path
from collections import Counter

# Add scripts/ and the C++ extension to the import path
_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_ROOT / "scripts"))
sys.path.insert(0, str(_ROOT / "nciipc-prep" / "build"))


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

_OCR_MIN_CHARS_PER_PAGE = 40  # below this average, a page is probably a scanned image


def _try_ocr(path: str) -> str | None:
    """
    OCR fallback for scanned/image PDFs where pdfplumber's text layer is
    empty or near-empty. Requires pytesseract + pdf2image (pip) plus the
    system tesseract-ocr and poppler binaries — all optional; returns None
    if any of them is missing so the caller just keeps whatever pdfplumber
    already extracted rather than failing outright.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        print("  OCR fallback unavailable: pip install pytesseract pdf2image "
              "(plus the system tesseract-ocr and poppler binaries) to enable it.")
        return None

    try:
        images = convert_from_path(path)
        text = "\n".join(pytesseract.image_to_string(img) for img in images)
        return text if text.strip() else None
    except Exception as exc:
        print(f"  OCR attempt failed ({type(exc).__name__}: {exc}) — "
              f"check that tesseract-ocr and poppler are installed.")
        return None


def extract_from_pdf(path: str) -> str:
    import pdfplumber
    text = ""
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"

    # A likely scanned/image PDF: pdfplumber's text layer found almost nothing.
    if page_count and (len(text) / page_count) < _OCR_MIN_CHARS_PER_PAGE:
        print(f"  Only {len(text)} chars across {page_count} page(s) — looks scanned, trying OCR ...")
        ocr_text = _try_ocr(path)
        if ocr_text:
            text = ocr_text

    # Drop repeated lines (headers / footers / page numbers)
    lines = text.splitlines()
    freq = Counter(l.strip() for l in lines)
    return "\n".join(l for l in lines if freq[l.strip()] <= 5)


def extract_from_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_from_url(url: str) -> str:
    import trafilatura
    downloaded = trafilatura.fetch_url(url)
    text = trafilatura.extract(downloaded, include_tables=True)
    if not text:
        raise ValueError(f"Could not extract readable content from: {url}")
    return text


def extract_text(source: str) -> str:
    if source.startswith("http://") or source.startswith("https://"):
        print(f"  Fetching URL ...")
        return extract_from_url(source)
    suffix = Path(source).suffix.lower()
    if suffix == ".pdf":
        print(f"  Extracting PDF ...")
        return extract_from_pdf(source)
    if suffix == ".docx":
        print(f"  Extracting DOCX ...")
        return extract_from_docx(source)
    if suffix in (".txt", ".text"):
        print(f"  Reading TXT ...")
        return Path(source).read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .docx, .txt, or a URL.")


# ---------------------------------------------------------------------------
# Full pipeline
# ---------------------------------------------------------------------------

def run(source: str) -> dict:
    """
    End-to-end: source (PDF path or URL) -> scores + LLM findings dict.

    Returns the score_documents() dict with an added "finding" key per control.
    LLM errors are caught gracefully; finding is set to "" if unavailable.
    """
    from scorer import score_documents
    from explain import explain_control

    # Step 1: extract
    print("[1/3] Extracting text ...")
    text = extract_text(source)
    words = len(text.split())
    if words < 50:
        raise ValueError(f"Too little text extracted ({words} words). Check the source.")
    print(f"      {words:,} words extracted.")

    # Step 2: BM25 + LF scoring
    print("[2/3] Scoring ...")
    result = score_documents([text])
    print(f"      Company score: {result['company_score']:.1f}/100")

    # Step 3: LLM findings + remediation
    print("[3/3] Generating findings (OpenRouter) ...")
    for control in ("PC1", "PC2", "PC3"):
        data = result[control]
        try:
            explanation = explain_control(
                control, data["score"], data["maturity"], data["fields"]
            )
            finding, remediation = explanation["finding"], explanation["remediation"]
        except Exception as e:
            finding, remediation = "", []
            print(f"      LLM unavailable for {control}: {e}")
        result[control]["finding"] = finding
        result[control]["remediation"] = remediation
        print(f"      {control}: {data['score']:.1f}/100  ({data['maturity']})")

    return result


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    if len(sys.argv) < 2:
        print("Usage: python scripts/pipeline.py <pdf_path_or_url>")
        sys.exit(1)

    source = sys.argv[1]
    print(f"\nNCIIPC Compliance Pipeline")
    print(f"Source: {source}\n")

    result = run(source)

    print("\n" + "=" * 52)
    print(f"  Company Score : {result['company_score']:.1f} / 100")
    print("=" * 52)
    for ctrl in ("PC1", "PC2", "PC3"):
        d = result[ctrl]
        print(f"  {ctrl}  {d['score']:>5.1f}/100  {d['maturity']}")
        if d.get("finding"):
            print(f"       {d['finding'][:180]}")
        for action in d.get("remediation", []):
            print(f"       -> {action[:160]}")
    print("=" * 52)

    out = _ROOT / "data" / "processed" / "scores.json"
    out.parent.mkdir(parents=True, exist_ok=True)
    with open(out, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2)
    print(f"\nFull results saved to {out}")
